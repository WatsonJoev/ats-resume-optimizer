"""
Multi-Agent Orchestrator
=========================
Coordinates the ATS Evaluation, Resume Handling, and Master agents
for complete resume optimization workflows.
"""

from typing import Dict, List, Optional, Any
from pathlib import Path
import json
import hashlib
from datetime import datetime

from agents import ATSEvaluationAgent, ResumeHandlingAgent, MasterAgent
from config import get_settings


class ATSOrchestrator:
    """
    Orchestrates the multi-agent system for ATS resume optimization.
    
    Flow:
    1. ATS Evaluation Agent analyzes JD and scores resume
    2. Resume Handling Agent optimizes the resume
    3. Master Agent reviews and approves changes
    """
    
    def __init__(self):
        """Initialize the orchestrator with all agents."""
        self.settings = get_settings()
        
        # Initialize agents
        print("[Orchestrator] Initializing agents...")
        self.ats_agent = ATSEvaluationAgent()
        self.resume_agent = ResumeHandlingAgent()
        self.master_agent = MasterAgent()
        
        # Connect agents to master
        self.master_agent.set_agents(
            ats_agent=self.ats_agent,
            resume_agent=self.resume_agent
        )
        
        print("[Orchestrator] All agents initialized")
    
    def _generate_job_hash(self, jd_text: str, company: str = "") -> str:
        """Generate a unique hash for a job."""
        content = f"{company}:{jd_text[:500]}"
        return hashlib.md5(content.encode()).hexdigest()[:12]
    
    def _save_job_data(
        self,
        job_hash: str,
        jd_text: str,
        metadata: Dict[str, Any],
        result: Dict[str, Any]
    ) -> Path:
        """Save job data to local storage."""
        job_dir = self.settings.jobs_dir / job_hash
        job_dir.mkdir(parents=True, exist_ok=True)
        
        # Save job description
        job_data = {
            "hash": job_hash,
            "metadata": metadata,
            "jd_text": jd_text,
            "created_at": datetime.now().isoformat(),
        }
        
        with open(job_dir / "job.json", "w") as f:
            json.dump(job_data, f, indent=2)
        
        # Save result
        result_data = {
            **result,
            "saved_at": datetime.now().isoformat(),
        }
        
        # Find next version number
        existing_results = list(job_dir.glob("result_v*.json"))
        version = len(existing_results) + 1
        
        with open(job_dir / f"result_v{version}.json", "w") as f:
            json.dump(result_data, f, indent=2, default=str)
        
        # Save optimized resume in multiple formats if approved
        if result.get("decision") == "APPROVED" and result.get("final_resume"):
            resume_content = result["final_resume"]
            
            # Save as TXT
            with open(job_dir / f"resume_v{version}.txt", "w", encoding="utf-8") as f:
                f.write(resume_content)
            
            # Save as Markdown
            md_path = job_dir / f"resume_v{version}.md"
            md_path.write_text(resume_content, encoding="utf-8")
            
            # Save as DOCX (ATS-optimized format)
            self._save_as_docx(resume_content, job_dir / f"resume_v{version}.docx")
            
            print(f"[Orchestrator] Resume saved in TXT, MD, and DOCX formats")
        
        return job_dir
    
    def _save_as_docx(self, content: str, file_path: Path) -> bool:
        """
        Save resume content as ATS-optimized DOCX file.
        
        Args:
            content: Resume text content
            file_path: Path to save the DOCX file
            
        Returns:
            True if successful, False otherwise
        """
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            print("[Orchestrator] python-docx not installed, skipping DOCX export")
            return False
        
        try:
            doc = Document()
            
            # Set default font (ATS-safe: Calibri, Arial, or Times New Roman)
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Calibri'
            font.size = Pt(11)
            
            # Set margins (standard 1 inch)
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            lines = content.split('\n')
            
            for line in lines:
                line_stripped = line.strip()
                
                if not line_stripped:
                    doc.add_paragraph()
                    continue
                
                # Detect section headers
                is_header = self._is_section_header(line_stripped)
                
                if is_header:
                    p = doc.add_paragraph()
                    run = p.add_run(line_stripped.replace(':', '').strip().upper())
                    run.bold = True
                    run.font.size = Pt(12)
                    run.font.name = 'Calibri'
                elif line_stripped.startswith(('-', '•', '*', '·')):
                    # Bullet point
                    clean_line = line_stripped.lstrip('-•*· ').strip()
                    doc.add_paragraph(clean_line, style='List Bullet')
                elif ':' in line_stripped and line_stripped.index(':') < 20:
                    # Key-value pair (like "Employer: Company Name")
                    p = doc.add_paragraph()
                    parts = line_stripped.split(':', 1)
                    run = p.add_run(parts[0] + ':')
                    run.bold = True
                    if len(parts) > 1:
                        p.add_run(' ' + parts[1].strip())
                else:
                    doc.add_paragraph(line_stripped)
            
            doc.save(str(file_path))
            return True
            
        except Exception as e:
            print(f"[Orchestrator] Error saving DOCX: {e}")
            return False
    
    def _is_section_header(self, line: str) -> bool:
        """Check if a line is a section header."""
        line_upper = line.upper().strip()
        
        # Common resume section headers
        headers = [
            'OVERVIEW', 'SUMMARY', 'OBJECTIVE', 'PROFILE',
            'EXPERIENCE', 'WORK EXPERIENCE', 'PROFESSIONAL EXPERIENCE', 'EMPLOYMENT',
            'EDUCATION', 'ACADEMIC', 'QUALIFICATIONS',
            'SKILLS', 'TECHNICAL SKILLS', 'CORE COMPETENCIES', 'EXPERTISE',
            'CERTIFICATIONS', 'CERTIFICATES', 'LICENSES',
            'PROJECTS', 'KEY PROJECTS', 'PORTFOLIO',
            'AWARDS', 'HONORS', 'ACHIEVEMENTS',
            'PROFESSIONAL SUMMARY', 'CAREER SUMMARY',
            'RESPONSIBILITIES', 'KEY RESPONSIBILITIES'
        ]
        
        # Check if line matches a known header
        for header in headers:
            if header in line_upper or line_upper.replace(':', '').strip() == header:
                return True
        
        # Check if ALL CAPS and reasonable length
        if line.isupper() and 3 < len(line) < 40:
            return True
        
        return False
    
    def optimize(
        self,
        jd_text: str,
        resume_text: str,
        user_skills: List[str] = None,
        job_metadata: Dict[str, str] = None,
        mode: str = "balanced",
        max_iterations: int = 3,
        save_results: bool = True
    ) -> Dict[str, Any]:
        """
        Run the complete optimization workflow.
        
        Args:
            jd_text: Job description text
            resume_text: Resume text
            user_skills: List of user's confirmed skills
            job_metadata: Job metadata (title, company, location, url)
            mode: Optimization mode (conservative, balanced, aggressive)
            max_iterations: Maximum review iterations
            save_results: Whether to save results locally
            
        Returns:
            Dict with optimization results
        """
        job_metadata = job_metadata or {}
        user_skills = user_skills or []
        
        print("\n" + "="*60)
        print("ATS RESUME OPTIMIZER - Starting Optimization")
        print("="*60)
        
        if job_metadata:
            print(f"Job: {job_metadata.get('title', 'Unknown')} at {job_metadata.get('company', 'Unknown')}")
        print(f"Mode: {mode}")
        print(f"Max iterations: {max_iterations}")
        print("="*60 + "\n")
        
        # Run the master orchestration
        result = self.master_agent.orchestrate_optimization(
            resume_text=resume_text,
            jd_text=jd_text,
            user_skills=user_skills,
            mode=mode,
            max_iterations=max_iterations
        )
        
        # Add metadata to result
        result["job_metadata"] = job_metadata
        result["mode"] = mode
        result["user_skills"] = user_skills
        
        # Save results if requested
        if save_results:
            job_hash = self._generate_job_hash(
                jd_text, 
                job_metadata.get("company", "")
            )
            result["job_hash"] = job_hash
            
            save_path = self._save_job_data(
                job_hash, jd_text, job_metadata, result
            )
            result["saved_to"] = str(save_path)
            print(f"\n[Orchestrator] Results saved to: {save_path}")
        
        print("\n" + "="*60)
        print(f"OPTIMIZATION COMPLETE - Decision: {result['decision']}")
        print(f"Iterations: {result['iterations']}")
        print("="*60 + "\n")
        
        return result
    
    def quick_score(
        self,
        jd_text: str,
        resume_text: str
    ) -> Dict[str, Any]:
        """
        Get a quick ATS score without full optimization.
        
        Args:
            jd_text: Job description text
            resume_text: Resume text
            
        Returns:
            Dict with quick score results
        """
        print("[Orchestrator] Running quick ATS score...")
        
        score_result = self.ats_agent.quick_score(jd_text, resume_text)
        
        return {
            "type": "quick_score",
            "result": score_result,
        }
    
    def evaluate_only(
        self,
        jd_text: str,
        resume_text: str,
        user_skills: List[str] = None
    ) -> Dict[str, Any]:
        """
        Run only ATS evaluation without optimization.
        
        Args:
            jd_text: Job description text
            resume_text: Resume text
            user_skills: User's skills
            
        Returns:
            Dict with evaluation results
        """
        print("[Orchestrator] Running ATS evaluation only...")
        
        return self.ats_agent.evaluate(jd_text, resume_text, user_skills)
    
    def list_saved_jobs(self) -> List[Dict[str, Any]]:
        """List all saved job optimizations."""
        jobs = []
        
        if not self.settings.jobs_dir.exists():
            return jobs
        
        for job_dir in self.settings.jobs_dir.iterdir():
            if job_dir.is_dir():
                job_file = job_dir / "job.json"
                if job_file.exists():
                    with open(job_file) as f:
                        job_data = json.load(f)
                    
                    # Count versions
                    versions = len(list(job_dir.glob("result_v*.json")))
                    job_data["versions"] = versions
                    jobs.append(job_data)
        
        # Sort by created_at descending
        jobs.sort(key=lambda x: x.get("created_at", ""), reverse=True)
        
        return jobs
    
    def get_job_result(self, job_hash: str, version: int = None) -> Optional[Dict[str, Any]]:
        """
        Get a saved job result.
        
        Args:
            job_hash: Job hash
            version: Specific version (latest if None)
            
        Returns:
            Job result or None
        """
        job_dir = self.settings.jobs_dir / job_hash
        
        if not job_dir.exists():
            return None
        
        if version is None:
            # Get latest version
            results = sorted(job_dir.glob("result_v*.json"))
            if not results:
                return None
            result_file = results[-1]
        else:
            result_file = job_dir / f"result_v{version}.json"
        
        if not result_file.exists():
            return None
        
        with open(result_file) as f:
            return json.load(f)


# Global orchestrator instance
_orchestrator: Optional[ATSOrchestrator] = None


def get_orchestrator() -> ATSOrchestrator:
    """Get the global orchestrator instance."""
    global _orchestrator
    if _orchestrator is None:
        _orchestrator = ATSOrchestrator()
    return _orchestrator
