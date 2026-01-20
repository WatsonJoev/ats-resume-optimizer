"""
Resume Parser Service
======================
Parses resumes from various formats (DOCX, PDF, TXT) into structured text.
"""

import re
from pathlib import Path
from typing import Dict, List, Optional, Any
from dataclasses import dataclass, field


@dataclass
class ParsedResume:
    """Structured resume data."""
    raw_text: str
    sections: Dict[str, str] = field(default_factory=dict)
    contact_info: Dict[str, str] = field(default_factory=dict)
    skills: List[str] = field(default_factory=list)
    file_path: Optional[str] = None
    format: str = "txt"


class ResumeParser:
    """
    Parser for extracting text and structure from resumes.
    
    Supports:
    - Plain text (.txt)
    - Microsoft Word (.docx)
    - PDF (.pdf) - requires pdfplumber
    """
    
    # Section header patterns
    SECTION_PATTERNS = {
        "summary": r"(?:summary|profile|objective|about\s*me)",
        "experience": r"(?:experience|work\s*history|employment|professional\s*experience)",
        "education": r"(?:education|academic|qualifications|degrees?)",
        "skills": r"(?:skills|technical\s*skills|competencies|technologies|expertise)",
        "projects": r"(?:projects|portfolio|work\s*samples)",
        "certifications": r"(?:certifications?|certificates?|licenses?)",
        "awards": r"(?:awards?|honors?|achievements?)",
    }
    
    def __init__(self):
        """Initialize the resume parser."""
        pass
    
    def parse_file(self, file_path: str) -> ParsedResume:
        """
        Parse a resume file.
        
        Args:
            file_path: Path to the resume file
            
        Returns:
            ParsedResume with extracted content
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"Resume file not found: {file_path}")
        
        suffix = path.suffix.lower()
        
        if suffix == ".txt":
            text = self._parse_txt(path)
        elif suffix == ".docx":
            text = self._parse_docx(path)
        elif suffix == ".pdf":
            text = self._parse_pdf(path)
        else:
            raise ValueError(f"Unsupported file format: {suffix}")
        
        return self._structure_resume(text, str(path), suffix[1:])
    
    def parse_text(self, text: str) -> ParsedResume:
        """
        Parse resume from raw text.
        
        Args:
            text: Raw resume text
            
        Returns:
            ParsedResume with extracted content
        """
        return self._structure_resume(text, None, "txt")
    
    def _parse_txt(self, path: Path) -> str:
        """Parse plain text file."""
        return path.read_text(encoding="utf-8")
    
    def _parse_docx(self, path: Path) -> str:
        """Parse Microsoft Word document."""
        try:
            from docx import Document
        except ImportError:
            raise ImportError(
                "python-docx is required for DOCX parsing. "
                "Install with: pip install python-docx"
            )
        
        doc = Document(path)
        paragraphs = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    paragraphs.append(row_text)
        
        return "\n".join(paragraphs)
    
    def _parse_pdf(self, path: Path) -> str:
        """Parse PDF document."""
        try:
            import pdfplumber
        except ImportError:
            raise ImportError(
                "pdfplumber is required for PDF parsing. "
                "Install with: pip install pdfplumber"
            )
        
        text_parts = []
        
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        
        return "\n".join(text_parts)
    
    def _structure_resume(
        self, 
        text: str, 
        file_path: Optional[str],
        format: str
    ) -> ParsedResume:
        """
        Extract structure from resume text.
        
        Args:
            text: Raw resume text
            file_path: Original file path
            format: File format
            
        Returns:
            ParsedResume with sections and metadata
        """
        sections = self._extract_sections(text)
        contact_info = self._extract_contact_info(text)
        skills = self._extract_skills(text)
        
        return ParsedResume(
            raw_text=text,
            sections=sections,
            contact_info=contact_info,
            skills=skills,
            file_path=file_path,
            format=format,
        )
    
    def _extract_sections(self, text: str) -> Dict[str, str]:
        """
        Extract named sections from resume text.
        
        Args:
            text: Resume text
            
        Returns:
            Dict mapping section names to content
        """
        sections = {}
        lines = text.split('\n')
        
        current_section = "header"
        current_content = []
        
        for line in lines:
            line_lower = line.lower().strip()
            
            # Check if this line is a section header
            found_section = None
            for section_name, pattern in self.SECTION_PATTERNS.items():
                if re.search(f"^{pattern}\\s*:?\\s*$", line_lower):
                    found_section = section_name
                    break
                # Also check for headers without trailing content
                if re.search(f"^{pattern}$", line_lower):
                    found_section = section_name
                    break
            
            if found_section:
                # Save previous section
                if current_content:
                    sections[current_section] = "\n".join(current_content).strip()
                current_section = found_section
                current_content = []
            else:
                current_content.append(line)
        
        # Save last section
        if current_content:
            sections[current_section] = "\n".join(current_content).strip()
        
        return sections
    
    def _extract_contact_info(self, text: str) -> Dict[str, str]:
        """
        Extract contact information from resume.
        
        Args:
            text: Resume text
            
        Returns:
            Dict with contact details
        """
        contact = {}
        
        # Email
        email_match = re.search(r'\b[\w.-]+@[\w.-]+\.\w+\b', text)
        if email_match:
            contact["email"] = email_match.group()
        
        # Phone (various formats)
        phone_patterns = [
            r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b',  # US format
            r'\+\d{1,3}[-.\s]?\d{1,4}[-.\s]?\d{1,4}[-.\s]?\d{1,9}',  # International
        ]
        for pattern in phone_patterns:
            phone_match = re.search(pattern, text)
            if phone_match:
                contact["phone"] = phone_match.group()
                break
        
        # LinkedIn
        linkedin_match = re.search(
            r'(?:linkedin\.com/in/|linkedin:\s*)([a-zA-Z0-9-]+)',
            text,
            re.IGNORECASE
        )
        if linkedin_match:
            contact["linkedin"] = linkedin_match.group(1)
        
        # GitHub
        github_match = re.search(
            r'(?:github\.com/|github:\s*)([a-zA-Z0-9-]+)',
            text,
            re.IGNORECASE
        )
        if github_match:
            contact["github"] = github_match.group(1)
        
        # Location (city, state pattern)
        location_match = re.search(
            r'\b([A-Z][a-z]+(?:\s+[A-Z][a-z]+)?),?\s*([A-Z]{2})\b',
            text
        )
        if location_match:
            contact["location"] = f"{location_match.group(1)}, {location_match.group(2)}"
        
        return contact
    
    def _extract_skills(self, text: str) -> List[str]:
        """
        Extract skills from resume.
        
        Args:
            text: Resume text
            
        Returns:
            List of identified skills
        """
        skills = set()
        text_lower = text.lower()
        
        # Common technical skills to look for
        tech_skills = [
            "python", "java", "javascript", "typescript", "c++", "c#", "go", "rust",
            "ruby", "php", "swift", "kotlin", "scala", "r", "matlab",
            "react", "angular", "vue", "node.js", "express", "django", "flask",
            "fastapi", "spring", "rails", ".net", "laravel",
            "aws", "azure", "gcp", "docker", "kubernetes", "terraform",
            "jenkins", "gitlab", "github actions", "circleci",
            "sql", "postgresql", "mysql", "mongodb", "redis", "elasticsearch",
            "dynamodb", "cassandra", "oracle", "sqlite",
            "git", "linux", "bash", "powershell",
            "html", "css", "sass", "tailwind", "bootstrap",
            "rest api", "graphql", "grpc", "websocket",
            "machine learning", "deep learning", "nlp", "computer vision",
            "tensorflow", "pytorch", "keras", "scikit-learn",
            "pandas", "numpy", "matplotlib", "jupyter",
            "agile", "scrum", "kanban", "jira", "confluence",
        ]
        
        for skill in tech_skills:
            if skill in text_lower:
                skills.add(skill)
        
        # Look for skills in a skills section
        skills_section = re.search(
            r'(?:skills|technologies|competencies)[:\s]*\n(.*?)(?:\n\n|\n[A-Z])',
            text,
            re.IGNORECASE | re.DOTALL
        )
        
        if skills_section:
            # Parse comma or bullet separated skills
            skills_text = skills_section.group(1)
            # Split by common delimiters
            skill_items = re.split(r'[,•·|\n]', skills_text)
            for item in skill_items:
                item = item.strip().strip('-').strip()
                if item and len(item) < 50:  # Reasonable skill length
                    skills.add(item.lower())
        
        return sorted(list(skills))
    
    def to_dict(self, parsed: ParsedResume) -> Dict[str, Any]:
        """Convert ParsedResume to dictionary for JSON serialization."""
        return {
            "raw_text": parsed.raw_text,
            "sections": parsed.sections,
            "contact_info": parsed.contact_info,
            "skills": parsed.skills,
            "file_path": parsed.file_path,
            "format": parsed.format,
        }
