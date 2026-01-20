"""
Local Storage Manager
======================
Handles local storage of resumes, job data, and user settings.
"""

import json
import shutil
from pathlib import Path
from typing import Dict, List, Optional, Any
from datetime import datetime

import sys
sys.path.insert(0, str(Path(__file__).parent.parent))

from config import get_settings


class LocalStorage:
    """
    Manages local storage for the ATS Resume Optimizer.
    
    Storage structure:
    ~/.ats-extension/ (or configured data_dir)
    ├── baseline/
    │   └── resume.txt (or .docx, .pdf)
    ├── jobs/
    │   └── {job_hash}/
    │       ├── job.json
    │       ├── result_v1.json
    │       ├── resume_v1.txt
    │       └── ...
    └── settings.json
    """
    
    def __init__(self):
        """Initialize local storage."""
        self.settings = get_settings()
        self._ensure_directories()
    
    def _ensure_directories(self):
        """Ensure all required directories exist."""
        self.settings.data_dir.mkdir(parents=True, exist_ok=True)
        self.settings.baseline_dir.mkdir(parents=True, exist_ok=True)
        self.settings.jobs_dir.mkdir(parents=True, exist_ok=True)
    
    # ==================== Baseline Resume ====================
    
    def save_baseline_resume(self, content: str, filename: str = "resume.txt") -> Path:
        """
        Save the user's baseline resume.
        
        Args:
            content: Resume content
            filename: Filename to save as
            
        Returns:
            Path to saved file
        """
        # Backup existing baseline if present
        existing = list(self.settings.baseline_dir.glob("resume.*"))
        if existing:
            backup_dir = self.settings.baseline_dir / "backups"
            backup_dir.mkdir(exist_ok=True)
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            for f in existing:
                shutil.copy(f, backup_dir / f"{f.stem}_{timestamp}{f.suffix}")
        
        # Save new baseline
        file_path = self.settings.baseline_dir / filename
        file_path.write_text(content, encoding="utf-8")
        
        return file_path
    
    def get_baseline_resume(self) -> Optional[str]:
        """
        Get the user's baseline resume content.
        Supports .txt, .md, .docx, and .pdf files.
        
        Returns:
            Resume content or None if not found
        """
        baseline_path = self.get_baseline_path()
        if not baseline_path:
            return None
        
        return self._read_resume_file(baseline_path)
    
    def _read_resume_file(self, file_path: Path) -> Optional[str]:
        """
        Read resume content from various file formats.
        
        Args:
            file_path: Path to the resume file
            
        Returns:
            Resume text content or None
        """
        suffix = file_path.suffix.lower()
        
        try:
            if suffix in [".txt", ".md"]:
                return file_path.read_text(encoding="utf-8")
            
            elif suffix == ".docx":
                return self._parse_docx(file_path)
            
            elif suffix == ".pdf":
                return self._parse_pdf(file_path)
            
            else:
                return None
        except Exception as e:
            print(f"Error reading resume file {file_path}: {e}")
            return None
    
    def _parse_docx(self, file_path: Path) -> Optional[str]:
        """Parse Microsoft Word document."""
        try:
            from docx import Document
        except ImportError:
            print("python-docx not installed. Install with: pip install python-docx")
            return None
        
        try:
            doc = Document(file_path)
            paragraphs = []
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        paragraphs.append(row_text)
            
            return "\n".join(paragraphs)
        except Exception as e:
            print(f"Error parsing DOCX: {e}")
            return None
    
    def _parse_pdf(self, file_path: Path) -> Optional[str]:
        """Parse PDF document."""
        try:
            import pdfplumber
        except ImportError:
            print("pdfplumber not installed. Install with: pip install pdfplumber")
            return None
        
        try:
            text_parts = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            return "\n".join(text_parts)
        except Exception as e:
            print(f"Error parsing PDF: {e}")
            return None
    
    def get_baseline_path(self) -> Optional[Path]:
        """
        Get path to baseline resume file.
        Searches for any resume file in the baseline directory.
        
        Returns:
            Path to the first found resume file or None
        """
        # First check for files named "resume.*"
        for ext in [".txt", ".md", ".docx", ".pdf"]:
            file_path = self.settings.baseline_dir / f"resume{ext}"
            if file_path.exists():
                return file_path
        
        # Then check for any .docx or .pdf file in the baseline folder
        for ext in ["*.docx", "*.pdf", "*.txt", "*.md"]:
            files = list(self.settings.baseline_dir.glob(ext))
            if files:
                # Return the most recently modified file
                files.sort(key=lambda x: x.stat().st_mtime, reverse=True)
                return files[0]
        
        return None
    
    def get_all_baseline_files(self) -> List[Path]:
        """
        Get all resume files in the baseline directory.
        
        Returns:
            List of paths to resume files
        """
        files = []
        for ext in ["*.docx", "*.pdf", "*.txt", "*.md"]:
            files.extend(self.settings.baseline_dir.glob(ext))
        
        # Sort by modification time (newest first)
        files.sort(key=lambda x: x.stat().st_mtime, reverse=True)
        return files
    
    def has_baseline_resume(self) -> bool:
        """Check if a baseline resume exists."""
        return self.get_baseline_path() is not None
    
    def parse_uploaded_file(self, uploaded_file) -> Optional[str]:
        """
        Parse an uploaded file (from Streamlit file_uploader).
        
        Args:
            uploaded_file: Streamlit UploadedFile object
            
        Returns:
            Resume text content or None
        """
        import tempfile
        
        suffix = Path(uploaded_file.name).suffix.lower()
        
        try:
            if suffix in [".txt", ".md"]:
                return uploaded_file.read().decode("utf-8")
            
            elif suffix == ".docx":
                # Save to temp file and parse
                with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                    tmp.write(uploaded_file.read())
                    tmp_path = Path(tmp.name)
                
                result = self._parse_docx(tmp_path)
                tmp_path.unlink()  # Clean up temp file
                return result
            
            elif suffix == ".pdf":
                # Save to temp file and parse
                with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                    tmp.write(uploaded_file.read())
                    tmp_path = Path(tmp.name)
                
                result = self._parse_pdf(tmp_path)
                tmp_path.unlink()  # Clean up temp file
                return result
            
            else:
                return None
        except Exception as e:
            print(f"Error parsing uploaded file: {e}")
            return None
    
    # ==================== User Skills ====================
    
    def save_user_skills(self, skills: List[str]) -> Path:
        """
        Save user's skills list.
        
        Args:
            skills: List of skills
            
        Returns:
            Path to saved file
        """
        file_path = self.settings.baseline_dir / "skills.json"
        
        data = {
            "skills": skills,
            "updated_at": datetime.now().isoformat(),
        }
        
        with open(file_path, "w") as f:
            json.dump(data, f, indent=2)
        
        return file_path
    
    def get_user_skills(self) -> List[str]:
        """
        Get user's saved skills.
        
        Returns:
            List of skills
        """
        file_path = self.settings.baseline_dir / "skills.json"
        
        if not file_path.exists():
            return []
        
        with open(file_path) as f:
            data = json.load(f)
        
        return data.get("skills", [])
    
    # ==================== Job Data ====================
    
    def list_jobs(self) -> List[Dict[str, Any]]:
        """
        List all saved jobs.
        
        Returns:
            List of job summaries
        """
        jobs = []
        
        if not self.settings.jobs_dir.exists():
            return jobs
        
        for job_dir in self.settings.jobs_dir.iterdir():
            if job_dir.is_dir():
                job_file = job_dir / "job.json"
                if job_file.exists():
                    with open(job_file) as f:
                        job_data = json.load(f)
                    
                    # Add version count
                    versions = len(list(job_dir.glob("result_v*.json")))
                    job_data["versions"] = versions
                    job_data["path"] = str(job_dir)
                    jobs.append(job_data)
        
        # Sort by created_at descending
        jobs.sort(key=lambda x: x.get("created_at", ""), reverse=True)
        
        return jobs
    
    def get_job(self, job_hash: str) -> Optional[Dict[str, Any]]:
        """
        Get job data by hash.
        
        Args:
            job_hash: Job hash
            
        Returns:
            Job data or None
        """
        job_dir = self.settings.jobs_dir / job_hash
        job_file = job_dir / "job.json"
        
        if not job_file.exists():
            return None
        
        with open(job_file) as f:
            return json.load(f)
    
    def get_job_result(self, job_hash: str, version: int = None) -> Optional[Dict[str, Any]]:
        """
        Get job optimization result.
        
        Args:
            job_hash: Job hash
            version: Specific version (latest if None)
            
        Returns:
            Result data or None
        """
        job_dir = self.settings.jobs_dir / job_hash
        
        if not job_dir.exists():
            return None
        
        if version is None:
            # Get latest version
            results = sorted(job_dir.glob("result_v*.json"))
            if not results:
                return None
            result_file = results[-1]
        else:
            result_file = job_dir / f"result_v{version}.json"
        
        if not result_file.exists():
            return None
        
        with open(result_file) as f:
            return json.load(f)
    
    def get_optimized_resume(self, job_hash: str, version: int = None) -> Optional[str]:
        """
        Get optimized resume for a job.
        
        Args:
            job_hash: Job hash
            version: Specific version (latest if None)
            
        Returns:
            Resume content or None
        """
        job_dir = self.settings.jobs_dir / job_hash
        
        if not job_dir.exists():
            return None
        
        if version is None:
            # Get latest version
            resumes = sorted(job_dir.glob("resume_v*.txt"))
            if not resumes:
                return None
            resume_file = resumes[-1]
        else:
            resume_file = job_dir / f"resume_v{version}.txt"
        
        if not resume_file.exists():
            return None
        
        return resume_file.read_text(encoding="utf-8")
    
    def delete_job(self, job_hash: str) -> bool:
        """
        Delete a job and all its data.
        
        Args:
            job_hash: Job hash
            
        Returns:
            True if deleted
        """
        job_dir = self.settings.jobs_dir / job_hash
        
        if not job_dir.exists():
            return False
        
        shutil.rmtree(job_dir)
        return True
    
    # ==================== User Settings ====================
    
    def save_settings(self, settings: Dict[str, Any]) -> Path:
        """
        Save user settings.
        
        Args:
            settings: Settings dict
            
        Returns:
            Path to saved file
        """
        file_path = self.settings.data_dir / "user_settings.json"
        
        settings["updated_at"] = datetime.now().isoformat()
        
        with open(file_path, "w") as f:
            json.dump(settings, f, indent=2)
        
        return file_path
    
    def get_user_settings(self) -> Dict[str, Any]:
        """
        Get user settings.
        
        Returns:
            Settings dict
        """
        file_path = self.settings.data_dir / "user_settings.json"
        
        if not file_path.exists():
            return {
                "default_mode": "balanced",
                "max_iterations": 3,
                "auto_save": True,
            }
        
        with open(file_path) as f:
            return json.load(f)
    
    # ==================== Resume Export ====================
    
    def save_resume_as_markdown(self, content: str, job_hash: str, version: int = None) -> Path:
        """
        Save optimized resume as Markdown file.
        
        Args:
            content: Resume text content
            job_hash: Job hash for organizing files
            version: Version number (auto-detected if None)
            
        Returns:
            Path to saved file
        """
        job_dir = self.settings.jobs_dir / job_hash
        job_dir.mkdir(parents=True, exist_ok=True)
        
        if version is None:
            existing = list(job_dir.glob("resume_v*.md"))
            version = len(existing) + 1
        
        file_path = job_dir / f"resume_v{version}.md"
        file_path.write_text(content, encoding="utf-8")
        
        return file_path
    
    def save_resume_as_docx(self, content: str, job_hash: str, version: int = None) -> Optional[Path]:
        """
        Save optimized resume as DOCX file with ATS-friendly formatting.
        
        Args:
            content: Resume text content (plain text or markdown)
            job_hash: Job hash for organizing files
            version: Version number (auto-detected if None)
            
        Returns:
            Path to saved file or None if failed
        """
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.style import WD_STYLE_TYPE
        except ImportError:
            print("python-docx not installed. Install with: pip install python-docx")
            return None
        
        job_dir = self.settings.jobs_dir / job_hash
        job_dir.mkdir(parents=True, exist_ok=True)
        
        if version is None:
            existing = list(job_dir.glob("resume_v*.docx"))
            version = len(existing) + 1
        
        file_path = job_dir / f"resume_v{version}.docx"
        
        # Create document with ATS-friendly formatting
        doc = Document()
        
        # Set default font (ATS-safe)
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Parse content and add to document
        lines = content.split('\n')
        current_section = None
        
        for line in lines:
            line = line.strip()
            if not line:
                # Add empty paragraph for spacing
                doc.add_paragraph()
                continue
            
            # Detect section headers (ALL CAPS or ends with :)
            is_header = (
                line.isupper() and len(line) > 3 and len(line) < 50
            ) or (
                line.endswith(':') and len(line) < 50
            ) or any(
                header in line.upper() for header in [
                    'OVERVIEW', 'SUMMARY', 'EXPERIENCE', 'EDUCATION', 
                    'SKILLS', 'CERTIFICATIONS', 'PROJECTS', 'WORK EXPERIENCE',
                    'PROFESSIONAL SUMMARY', 'TECHNICAL SKILLS', 'QUALIFICATIONS'
                ]
            )
            
            if is_header:
                # Add section header
                p = doc.add_paragraph()
                run = p.add_run(line.replace(':', '').strip())
                run.bold = True
                run.font.size = Pt(12)
                current_section = line
            elif line.startswith(('-', '•', '*', '·')):
                # Bullet point
                clean_line = line.lstrip('-•*· ').strip()
                p = doc.add_paragraph(clean_line, style='List Bullet')
            elif line.startswith(('Project #', 'Employer', 'Client', 'Period', 'Work mode', 'Title')):
                # Job details - keep formatting
                p = doc.add_paragraph()
                if ':' in line:
                    parts = line.split(':', 1)
                    run = p.add_run(parts[0] + ':')
                    run.bold = True
                    if len(parts) > 1:
                        p.add_run(parts[1])
                else:
                    p.add_run(line)
            else:
                # Regular paragraph
                p = doc.add_paragraph(line)
        
        # Save document
        doc.save(str(file_path))
        
        return file_path
    
    def export_optimized_resume(
        self, 
        job_hash: str, 
        formats: List[str] = None
    ) -> Dict[str, Optional[Path]]:
        """
        Export optimized resume in multiple formats.
        
        Args:
            job_hash: Job hash
            formats: List of formats to export ('txt', 'md', 'docx')
            
        Returns:
            Dict mapping format to file path
        """
        formats = formats or ['txt', 'md', 'docx']
        result = {}
        
        # Get the latest result
        job_result = self.get_job_result(job_hash)
        if not job_result:
            return result
        
        content = job_result.get('final_resume', '')
        if not content:
            return result
        
        # Find next version number
        job_dir = self.settings.jobs_dir / job_hash
        existing_txt = list(job_dir.glob("resume_v*.txt"))
        version = len(existing_txt) + 1
        
        # Export to each format
        if 'txt' in formats:
            txt_path = job_dir / f"resume_v{version}.txt"
            txt_path.write_text(content, encoding="utf-8")
            result['txt'] = txt_path
        
        if 'md' in formats:
            result['md'] = self.save_resume_as_markdown(content, job_hash, version)
        
        if 'docx' in formats:
            result['docx'] = self.save_resume_as_docx(content, job_hash, version)
        
        return result
    
    # ==================== Cleanup ====================
    
    def cleanup_old_jobs(self, days: int = 30) -> int:
        """
        Delete jobs older than specified days.
        
        Args:
            days: Age threshold in days
            
        Returns:
            Number of jobs deleted
        """
        from datetime import timedelta
        
        cutoff = datetime.now() - timedelta(days=days)
        deleted = 0
        
        for job in self.list_jobs():
            created_at = job.get("created_at", "")
            if created_at:
                try:
                    job_date = datetime.fromisoformat(created_at)
                    if job_date < cutoff:
                        self.delete_job(job["hash"])
                        deleted += 1
                except ValueError:
                    pass
        
        return deleted
